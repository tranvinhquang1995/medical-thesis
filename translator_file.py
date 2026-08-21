import os
import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from pypdf import PdfReader
from docx.shared import Pt, Inches

SYSTEM_INSTRUCTION_FILE = """
You are an expert medical translator specializing in translating professional medical documents, research papers, and academic theses between English and Vietnamese.
Your goal is to provide accurate, natural, and contextually appropriate translations.

Rules for Translation:
1. Contextual Translation: Translate full paragraphs or sentences as a whole. Do NOT translate word-for-word. Ensure natural phrasing, professional medical flow, and accurate scientific meaning.
2. Terminology Handling:
   - Common medical conditions and diseases with widely recognized Vietnamese names MUST be translated naturally into Vietnamese. Examples: "Heart failure" -> "Suy tim", "Diabetes mellitus" -> "Đái tháo đường", "Hypertension" -> "Tăng huyết áp", "Stroke" -> "Đột quỵ", "Myocardial infarction" -> "Nhồi máu cơ tim", "Pneumonia" -> "Viêm phổi".
   - Drug names (generic and brand names) MUST be kept exactly as they are in English (e.g., "Metformin", "Atorvastatin", "Aspirin", "Pembrolizumab"). Do NOT attempt to translate them.
   - Rare, highly specialized, or complex diseases, syndromes, or surgical procedures should remain in English (e.g., "Gilles de la Tourette syndrome").
   - Medical and professional abbreviations (e.g., "COPD", "CABG", "EF", "SLE", "CT-scan", "MRI", "PCR", "EGFR") MUST remain in English. Do NOT translate or expand them.
3. Stateful Coherence: You are translating a continuous document chunk-by-chunk. Maintain stylistic and terminological consistency with the previously translated content in this chat session.
4. Output: Return ONLY the translated text corresponding to the input chunk. Do not add conversational preambles, comments, or "Here is the translation:".
"""

def translate_docx(file_path: str, output_path: str, direction: str, api_key: str, model_name: str = "gemini-2.5-flash", progress_bar=None) -> str:
    """
    Translates a DOCX file while maintaining structure and layout.
    Uses a stateful Chat session for contextual coherence.
    """
    if not api_key:
        raise ValueError("API Key is required.")
        
    client = genai.Client(api_key=api_key)
    
    # Initialize stateful chat
    chat = client.chats.create(
        model=model_name,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION_FILE,
            temperature=0.3,
            max_output_tokens=8192,
        )
    )
    
    doc = Document(file_path)
    
    # Count total items for progress reporting
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    total_steps = total_paragraphs + sum(len(table.rows) * len(table.columns) for table in doc.tables)
    current_step = 0
    
    # Function to translate a text block using the stateful chat session
    def translate_block(text: str) -> str:
        if not text.strip():
            return text
        prompt = f"Translate the following chunk from {'English to Vietnamese' if direction == 'EN -> VI' else 'Vietnamese to English'}:\n\n{text}"
        try:
            response = chat.send_message(prompt)
            return response.text.strip()
        except Exception as e:
            st.error(f"Error in chat translation: {str(e)}")
            return text

    # Translate paragraphs (main body)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            # Save original run formatting if possible
            original_runs_format = []
            if paragraph.runs:
                first_run = paragraph.runs[0]
                bold = first_run.bold
                italic = first_run.italic
                font_name = first_run.font.name
                font_size = first_run.font.size
            else:
                bold, italic, font_name, font_size = False, False, None, None
                
            translated_text = translate_block(paragraph.text)
            
            # Clear text and write translated text
            paragraph.text = ""
            run = paragraph.add_run(translated_text)
            run.bold = bold
            run.italic = italic
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
                
        current_step += 1
        if progress_bar and total_steps > 0:
            progress_bar.progress(min(current_step / total_steps, 1.0), text=f"Dịch đoạn văn {i+1}/{total_paragraphs}...")

    # Translate tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Translate text in cell paragraphs
                for p in cell.paragraphs:
                    if p.text.strip():
                        translated_cell_text = translate_block(p.text)
                        p.text = translated_cell_text
                current_step += 1
                if progress_bar and total_steps > 0:
                    progress_bar.progress(min(current_step / total_steps, 1.0), text=f"Dịch bảng {t_idx+1}, ô ({r_idx},{c_idx})...")

    doc.save(output_path)
    return output_path

def translate_pdf(file_path: str, output_docx_path: str, direction: str, api_key: str, model_name: str = "gemini-2.5-flash", progress_bar=None) -> str:
    """
    Translates a PDF file by extracting its pages and generating a clean,
    well-formatted DOCX document with matching translation.
    Uses stateful Chat session.
    """
    if not api_key:
        raise ValueError("API Key is required.")
        
    client = genai.Client(api_key=api_key)
    
    # Initialize stateful chat
    chat = client.chats.create(
        model=model_name,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION_FILE,
            temperature=0.3,
            max_output_tokens=8192,
        )
    )
    
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)
    
    # Create a new Word document to store the translation
    doc = Document()
    doc.add_heading("BẢN DỊCH Y KHOA / MEDICAL TRANSLATION", level=1)
    doc.add_paragraph("Tài liệu được dịch tự động bằng ứng dụng Medical Thesis (Gemini 3.7 Flash).")
    doc.add_paragraph(f"Chiều dịch: {direction}")
    doc.add_page_break()
    
    for page_num in range(total_pages):
        page = reader.pages[page_num]
        text = page.extract_text()
        
        doc.add_heading(f"Trang {page_num + 1}", level=2)
        
        if text.strip():
            # Send text of the page as a single message to keep context page-by-page
            prompt = f"Translate the following page {page_num + 1} from {'English to Vietnamese' if direction == 'EN -> VI' else 'Vietnamese to English'}:\n\n{text}"
            try:
                response = chat.send_message(prompt)
                translated_text = response.text.strip()
                
                # Write paragraphs to DOCX
                for p_text in translated_text.split('\n\n'):
                    if p_text.strip():
                        doc.add_paragraph(p_text.strip())
            except Exception as e:
                st.error(f"Lỗi khi dịch trang {page_num + 1}: {str(e)}")
                doc.add_paragraph(f"[Lỗi dịch trang {page_num + 1}]")
        else:
            doc.add_paragraph("[Trang này không có chữ hoặc là trang quét ảnh]")
            
        if page_num < total_pages - 1:
            doc.add_page_break()
            
        if progress_bar:
            progress_bar.progress((page_num + 1) / total_pages, text=f"Đã dịch {page_num + 1}/{total_pages} trang PDF...")
            
    doc.save(output_docx_path)
    return output_docx_path
